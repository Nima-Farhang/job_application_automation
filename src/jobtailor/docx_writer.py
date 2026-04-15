from __future__ import annotations

from pathlib import Path

from docx import Document


def _clean_lines(text: str) -> list[str]:
    return [line.rstrip() for line in text.splitlines()]


def _extract_section(text: str, header: str, next_headers: list[str]) -> str:
    lines = _clean_lines(text)
    start_index = None
    for i, line in enumerate(lines):
        if line.strip() == header:
            start_index = i + 1
            break
    if start_index is None:
        return ""

    end_index = len(lines)
    for i in range(start_index, len(lines)):
        if lines[i].strip() in next_headers:
            end_index = i
            break
    return "\n".join(lines[start_index:end_index]).strip()


def save_simple_docx(text: str, output_path: Path) -> None:
    print(f"[docx] Writing DOCX: {output_path}")
    document = Document()
    for line in _clean_lines(text):
        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
            continue
        if stripped.startswith("• "):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(stripped[2:])
        else:
            document.add_paragraph(stripped)
    document.save(output_path)


def export_final_docs(stage3_text: str, output_dir: Path) -> tuple[Path, Path]:
    print(f"[docx] Exporting final documents into: {output_dir}")
    cv_text = _extract_section(
        stage3_text,
        "1️⃣ Final CV",
        ["2️⃣ Final Cover Letter", "3️⃣ Brief Strategic Evaluation"],
    )
    cover_letter_text = _extract_section(
        stage3_text,
        "2️⃣ Final Cover Letter",
        ["3️⃣ Brief Strategic Evaluation"],
    )

    if not cv_text:
        print("[docx] Final CV section not isolated; falling back to full Stage 3 output")
        cv_text = stage3_text
    if not cover_letter_text:
        print("[docx] Final cover letter section not isolated; using fallback text")
        cover_letter_text = "Cover letter could not be isolated from the Stage 3 output."

    cv_path = output_dir / "cv_final.docx"
    cover_path = output_dir / "cover_letter_final.docx"

    print(
        "[docx] Extracted sections: "
        f"cv_chars={len(cv_text)}, cover_letter_chars={len(cover_letter_text)}"
    )
    save_simple_docx(cv_text, cv_path)
    save_simple_docx(cover_letter_text, cover_path)
    return cv_path, cover_path
